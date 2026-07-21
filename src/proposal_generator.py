"""
proposal_generator.py
Renders an EstimateSummary into a client-ready Word (.docx) proposal:
- Cover section (client, date, prepared by, validity)
- Scope of work table (requirement, role, complexity)
- Effort & cost breakdown table
- Summary (subtotal, risk buffer, overhead, grand total)
- Terms & assumptions
"""

from datetime import datetime, timedelta
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.estimator import EstimateSummary
from src.config_manager import ProjectConfig
from src.utils import format_currency

HEADER_COLOR = RGBColor(0x1F, 0x4E, 0x79)


def _add_heading(doc: Document, text: str, size: int = 16, color: RGBColor = HEADER_COLOR):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _style_table_header(row):
    for cell in row.cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._tc.get_or_add_tcPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1F4E79")
        shading.append(shd)


def generate_proposal(
    summary: EstimateSummary,
    project_config: ProjectConfig,
    client_name: str,
    output_path: str,
) -> str:
    """Build the Word proposal document and save it to output_path. Returns the path."""
    doc = Document()

    # --- Cover section ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Project Effort & Cost Proposal")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = HEADER_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Prepared for {client_name}")
    sub_run.font.size = Pt(14)
    sub_run.italic = True

    doc.add_paragraph()

    today = datetime.now()
    valid_until = today + timedelta(days=project_config.proposal_validity_days)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta_rows = [
        ("Prepared by", project_config.prepared_by),
        ("Company", project_config.company_name),
        ("Date issued", today.strftime("%d %B %Y")),
        ("Valid until", valid_until.strftime("%d %B %Y")),
    ]
    for i, (label, value) in enumerate(meta_rows):
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        meta.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # --- Scope of work ---
    _add_heading(doc, "1. Scope of Work")
    scope_table = doc.add_table(rows=1, cols=4)
    scope_table.style = "Light Grid Accent 1"
    hdr = scope_table.rows[0].cells
    headers = ["Req. ID", "Description", "Role", "Complexity"]
    for cell, text in zip(hdr, headers):
        cell.text = text
    _style_table_header(scope_table.rows[0])

    for item in summary.line_items:
        row = scope_table.add_row().cells
        row[0].text = item.requirement_id
        row[1].text = item.description
        row[2].text = item.role
        row[3].text = item.complexity

    doc.add_paragraph()

    # --- Effort & cost breakdown ---
    _add_heading(doc, "2. Effort & Cost Breakdown")
    cost_table = doc.add_table(rows=1, cols=4)
    cost_table.style = "Light Grid Accent 1"
    hdr2 = cost_table.rows[0].cells
    headers2 = ["Req. ID", "Effort (person-days)", "Daily Rate", "Cost"]
    for cell, text in zip(hdr2, headers2):
        cell.text = text
    _style_table_header(cost_table.rows[0])

    for item in summary.line_items:
        row = cost_table.add_row().cells
        row[0].text = item.requirement_id
        row[1].text = f"{item.effort_days:.2f}"
        row[2].text = format_currency(item.daily_rate, summary.currency)
        row[3].text = format_currency(item.cost, summary.currency)

    doc.add_paragraph()

    # --- Summary ---
    _add_heading(doc, "3. Cost Summary")
    summary_table = doc.add_table(rows=5, cols=2)
    summary_rows = [
        ("Total effort (person-days)", f"{summary.subtotal_effort_days:.2f}"),
        ("Subtotal cost", format_currency(summary.subtotal_cost, summary.currency)),
        (
            f"Risk buffer ({summary.risk_buffer_percent:.0f}%)",
            format_currency(summary.risk_buffer_cost, summary.currency),
        ),
        (
            f"Overhead ({summary.overhead_percent:.0f}%)",
            format_currency(summary.overhead_cost, summary.currency),
        ),
        ("Grand Total", format_currency(summary.grand_total, summary.currency)),
    ]
    for i, (label, value) in enumerate(summary_rows):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        summary_table.rows[i].cells[1].text = value
        if label == "Grand Total":
            for cell in summary_table.rows[i].cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.bold = True
                        r.font.size = Pt(12)

    doc.add_paragraph()

    # --- Terms ---
    _add_heading(doc, "4. Terms & Assumptions", size=14)
    for term in project_config.proposal_terms:
        doc.add_paragraph(term, style="List Bullet")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
